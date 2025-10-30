# processor.py (backend logic) — improved variable detection
#
# Changes in this version:
# - Reordered top-level functions used by streamlit_app: configure, find_relevant_variables, fill_section_values, generate_doc_from_excel_map.
# - Improved find_relevant_variables/_prefill logic to more reliably detect two classes of changes:
#     * Standardized change: year/date patterns and FY patterns (e.g., "FY2023", "December 31st 2023", "Financial Year Ended 31 December 2023").
#     * Contextual change: key financial/benchmark figures (percentages, euro amounts, ratios, observation counts) when they appear near known labels such as "Net turnover", "Median", "Employees", "FTE", "Berry ratio", etc.
# - Adds a 'change_type' column (if not present) when pre-filling so the sheet can be used to review classification. This is non-breaking: other code still identifies core columns by name.
# - The LLM-based extraction (ask_variable_list) is still used and respected; regex-based detections augment LLM rows (no duplicates).
#
# Note: Behavior otherwise preserved. If you want this committed to the repo, I can prepare a branch and PR (I cannot push without your confirmation).

import requests
from docx import Document

from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.oxml.ns import qn


import pdfplumber
import pandas as pd

import os
import io
import time
import tempfile
import openpyxl

from collections import OrderedDict
from collections import defaultdict
from typing import Tuple, Dict, List, Optional
from pydantic import BaseModel, TypeAdapter

import re, math, json
from dataclasses import dataclass




# Azure OpenAI config (set by frontend)
API_KEY = None
API_ENDPOINT = None

class Row(BaseModel):
    variable_name: str
    old_value: str
    prompt: str
    new_value: str

    model_config = {
        "extra": "forbid",
        "str_strip_whitespace": True
    }

_row_list_adapter = TypeAdapter(List[Row])

def parse_rows_json(s: str) -> List[Row]:
    return _row_list_adapter.validate_json(s)


# -------------------------
# Top-level API (front-end)
# -------------------------

def configure(api_key: str, api_endpoint: str):
    """Called once by the Streamlit frontend to set API credentials/endpoints."""
    global API_KEY, API_ENDPOINT
    API_KEY = api_key
    API_ENDPOINT = api_endpoint


def find_relevant_variables(files: dict):
    """
    Used by Step 1 in the UI.
    - Build context from provided files
    - Optionally ask LLM to extract candidate placeholders
    - Augment LLM output with robust regex-based detection for:
        * Standardized changes (years/dates/FY)
        * Contextual changes around financial labels (percentages, euro amounts, counts)
    - Produce an annotated variables Excel and a fallback docx summarizing replacements.
    Returns (doc_path, filled_excel_path)
    """
    guidelines = files.get("guidelines") if files else None
    transcript = files.get("transcript") if files else None
    pdf = files.get("pdf") if files else None
    excel = files.get("excel") if files else None
    template = files.get("template") if files else None

    # Build context
    ctx = ""
    ctx += load_guidelines(guidelines)
    tr_text = load_transcript(transcript)
    if tr_text:
        ctx += ("\n\n" if ctx else "") + tr_text
    pdf_text = load_pdf(pdf)
    if pdf_text:
        ctx += ("\n\n" if ctx else "") + pdf_text

    template_text = load_template(template) if template else None
    if template_text:
        ctx += ("\n\n" if ctx else "") + template_text

    # 1) Use LLM to generate candidate rows (this also helps identify section-level prompts)
    maybe_prefilled_path = _prefill_last_year_from_prompts(excel, ctx)

    # 2) If LLM produced a prefilled workbook, use it, otherwise use original excel
    excel_for_processing = maybe_prefilled_path or excel

    # 3) Load-and-annotate replacements (this will call LLM to fill prompts in the workbook)
    replacements, filled_excel_path = load_and_annotate_replacements(excel_for_processing, ctx) if excel_for_processing else ({}, None)

    # 4) Augment the workbook with regex-based detections for standardized/contextual values
    try:
        filled_excel_path = _augment_with_regex_detections(excel_for_processing or filled_excel_path, ctx)
    except Exception:
        # Non-fatal: keep previously created filled_excel_path
        pass

    # Build a fallback doc summarizing replacements & context for preview
    doc_path = _build_fallback_docx(replacements, ctx)
    return (doc_path, filled_excel_path)


def fill_section_values(files):
    """
    Used by Step 2. Rebuilds context and fills section-level prompts via LLM.
    Returns path to annotated excel.
    """
    guidelines = files.get("guidelines") if files else None
    transcript = files.get("transcript") if files else None
    pdf = files.get("pdf") if files else None
    excel = files.get("excel") if files else None
    template = files.get("template") if files else None

    ctx = ""
    ctx += load_guidelines(guidelines)
    tr_text = load_transcript(transcript)
    if tr_text:
        ctx += ("\n\n" if ctx else "") + tr_text
    pdf_text = load_pdf(pdf)
    if pdf_text:
        ctx += ("\n\n" if ctx else "") + pdf_text
    template_text = load_template(template) if template else None
    if template_text:
        ctx += ("\n\n" if ctx else "") + template_text

    excel_for_processing = excel
    (rep, filled_path) = load_and_annotate_replacements(excel_for_processing, ctx)
    # Also run regex-only augmentation to capture standardized changes (years/dates)
    try:
        filled_path = _augment_with_regex_detections(excel_for_processing or filled_path, ctx)
    except Exception:
        pass
    return filled_path


def generate_doc_from_excel_map(file_map, context: str = ""):
    """
    Used by Step 3. Reads annotated excel, builds replacements, applies to template, returns (doc_path, filled_excel_path).
    """
    def _load_wb(x):
        try:
            if isinstance(x, str) and os.path.exists(x):
                return openpyxl.load_workbook(x)
            if hasattr(x, "read"):
                b = x.read()
                try:
                    x.seek(0)
                except Exception:
                    pass
                return openpyxl.load_workbook(io.BytesIO(b))
        except Exception:
            pass
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.append(["variable_name", "old_value", "prompt", "new_value"])
        return wb

    wb = _load_wb(file_map.get("excel"))
    ws = wb.active

    # --- header map (fallback to A..D) ---
    try:
        hdr = [str(c.value).strip().lower() if c.value else "" for c in ws[1]]
    except Exception:
        hdr = []
    def col(name, default_idx):
        return (hdr.index(name)+1) if name in hdr else default_idx
    c_old, c_new = col("old_value", 2), col("new_value", 4)

    repl = OrderedDict()
    for r in range(2, ws.max_row+1):
        old_s = (str(ws.cell(r, c_old).value).strip() if ws.cell(r, c_old).value else "")
        new_s = (str(ws.cell(r, c_new).value).strip() if ws.cell(r, c_new).value else "")
        if old_s and new_s and old_s != new_s:
            repl[old_s] = new_s

    # save annotated excel copy
    tmpdir = tempfile.mkdtemp()
    filled_excel_path = os.path.join(tmpdir, "variables_filled.xlsx")
    try:
        wb.save(filled_excel_path)
    except Exception:
        filled_excel_path = None

    # Try to apply replacements to template docx
    template = file_map.get("template")
    try:
        doc = Document(template) if template else None
    except Exception:
        doc = None

    if doc:
        if 'replace_first_page_placeholders_docx' in globals():
            replace_first_page_placeholders_docx(doc, repl)
        if 'replace_placeholders_docx' in globals():
            replace_placeholders_docx(doc, repl)
        out = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(out.name)
        return out.name, filled_excel_path

    if '_build_fallback_docx' in globals():
        return _build_fallback_docx(repl, context), filled_excel_path

    # ultimate fallback
    doc = Document(); doc.add_heading("Generated Document (Fallback)", 1)
    doc.add_paragraph("Template unavailable. Applied replacements:")
    for k, v in repl.items(): doc.add_paragraph(f"{k} → {v}")
    out = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(out.name)
    return out.name, filled_excel_path


# -------------------------
# Helpers used by top-level flows
# -------------------------

def ask_variable_list(context: str, wait_seconds: float = 2.0, max_retries: int = 6):
    """
    Ask LLM to return JSON list of Row objects. Retries until parseable.
    """
    headers = {"Content-Type": "application/json", "api-key": API_KEY}

    system_msg = (
        "You output ONLY JSON: an array of objects with exactly these keys: "
        "variable_name, old_value, prompt, new_value. No prose/markdown. "
        "Never output any object where a value equals its key name "
        '(e.g., "old_value":"old_value"). If you cannot find valid rows, return [].'
    )

    user_prompt = f"""
You are a meticulous placeholder auditor for corporate documents (transfer pricing + finance context).

OUTPUT
Return ONLY a JSON array. Each element must have exactly:
- "variable_name": string (for section-level use "SECTION:<name>")
- "old_value": string
- "prompt": string (non-empty ONLY if the placeholder is a whole section needing rewrite; otherwise "")
- "new_value": string (empty if <80% confident)

RULES
- No markdown, no comments, no code fences—JSON array only.
- Keep strings terse; prefer exact spans from CONTEXT.
- Identify changes in names/addresses/IDs; dates/fiscal years; financials; FAR; intercompany; benchmarks; governance; jurisdictions/regulations; and any sections to be rewritten.
- **old_value must be an exact substring of CONTEXT that is unique (occurs once).**
- **If a bare value would occur multiple times, expand old_value by including adjacent label/unit/words from CONTEXT until the span is unique (e.g., prefer 'employees: 14 FTEs' over '14').**
- **Do not invent header/example rows; do not use key names as values.**
- **For non-section placeholders, prompt must be empty (''); for SECTION:<name> rows, prompt must be non-empty.**
- **Deduplicate: do not output multiple rows with the same (variable_name, old_value).**

CONTEXT
<<<CONTEXT_START
{context}
CONTEXT_END>>>
""".strip()

    attempt = 0
    last_exception = None

    while attempt < max_retries:
        attempt += 1
        try:
            payload = {
                "messages": [
                    {"role": "system", "content": system_msg},
                    {"role": "user", "content": user_prompt},
                ],
                "temperature": 0,
                "top_p": 1,
                "seed": 7,
            }
            resp = requests.post(API_ENDPOINT, headers=headers, json=payload, timeout=60)
            resp.raise_for_status()
            data = resp.json()
            text = (
                data.get("choices", [{}])[0]
                    .get("message", {})
                    .get("content", "")
                    .strip()
            )
            rows = parse_rows_json(text)
            for r in rows:
                if r.prompt is None:
                    r.prompt = ""
                if r.new_value is None:
                    r.new_value = ""
            return rows
        except Exception as e:
            last_exception = e
            print(f"[ask_variable_list] attempt {attempt}/{max_retries} failed: {e}")
            time.sleep(wait_seconds)

    raise RuntimeError(f"ask_variable_list failed after {max_retries} attempts: {last_exception}")

def _classify_change_type(old_value: str) -> str:
    """
    Return 'Standardized change' for date/FY/year-like tokens, otherwise
    'Contextual change' for numbers/currency/percent-like tokens, else ''.
    """
    if not old_value:
        return ""
    try:
        s = str(old_value)
        if RE_FY.search(s) or RE_DATE_DMY.search(s) or RE_DATE_MDY.search(s) or RE_FINYEAR_PH.search(s) or RE_YR_RANGE.search(s) or RE_YEAR.search(s):
            return "Standardized change"
        if RE_PERCENT.search(s) or RE_CURRENCY.search(s) or RE_NUMBER.search(s):
            return "Contextual change"
    except Exception:
        pass
    return ""


def _detect_standard_and_context_spans(context: str):
    """
    Lightweight extractor that returns a list of tuples:
      (old_value: str, change_type: str, hint: str)
    This supplements LLM output with regex-based detections for:
      - FY/date/year tokens (Standardized change)
      - Labelled numbers near known labels (Contextual change)
    """
    out = []
    seen = set()

    if not context:
        return out

    # 1) Standard date/year/FY tokens (take unique matches)
    for rx in (RE_FY, RE_DATE_DMY, RE_DATE_MDY, RE_FINYEAR_PH, RE_FOR_YE_PH, RE_YR_RANGE, RE_YEAR):
        for m in rx.finditer(context):
            val = m.group(0).strip()
            if val and val not in seen:
                seen.add(val)
                out.append((val, "Standardized change", "financial_year"))

    # 2) Contextual: scan sentences for known labels and numbers
    for sent in sentence_splits(context):
        # only examine sentences that contain a label token
        if not LABELS_RE.search(sent):
            continue
        # find candidate numbers/currency/percents in the sentence
        for rx in (RE_PERCENT, RE_CURRENCY, RE_NUMBER):
            for m in rx.finditer(sent):
                val = canonicalize_num(m.group(0))
                if not val:
                    continue
                # avoid section numbers like "6.3"
                if RE_SECTIONNUM.match(val):
                    continue
                if val in seen:
                    continue
                seen.add(val)
                hint_m = LABELS_RE.search(sent)
                hint = hint_m.group(0) if hint_m else ""
                change_type = "Contextual change" if not RE_YEAR.fullmatch(val) else "Standardized change"
                out.append((val, change_type, hint or "labelled_value"))

    # 3) Return as list of tuples (old_value, change_type, hint)
    return out

def fill_excel_prompts(prompt: str, context: str, old_value: str, variable_name: str) -> str:
    """
    Call LLM to fill an individual Excel prompt cell.
    """
    headers = {"Content-Type": "application/json", "api-key": API_KEY}
    system_msg = (
        "You are an expert on Transfer Pricing and financial analysis. "
        "You are provided the old value of a variable, and the name of the variable."
        "Update this value with the latest data and STAY AS CLOSE TO OLD VALUE AS POSSIBLE"
        "Use either the internet or the context to infer the new value, and LEAVE IT THE SAME if it is not determinate. If you use the internet then cite sources. "
    )
    user_content = (
        f"old_value: {old_value}\n"
        f"variable_name: {variable_name}\n\n"
        f"Prompt: {prompt}\n\n"
        f"Context:\n{context}"
    )
    messages = [
        {"role": "system", "content": system_msg},
        {"role": "user", "content": user_content},
    ]
    resp = requests.post(API_ENDPOINT, headers=headers, json={"messages": messages})
    resp.raise_for_status()
    return resp.json()["choices"][0]["message"]["content"].strip()


def load_transcript(file) -> str:
    if not file:
        return ""
    try:
        doc = Document(file)
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception:
        return ""


def load_pdf(file) -> str:
    if not file:
        return ""
    pages, tables = [], []
    try:
        with pdfplumber.open(file) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                pages.append(f"--- Page {i} ---\n{text}")
                for table in page.extract_tables() or []:
                    try:
                        df = pd.DataFrame(table[1:], columns=table[0])
                        tables.append(f"--- Page {i} table ---\n" + df.to_csv(index=False))
                    except Exception:
                        continue
        return "\n\n".join(pages + tables)
    except Exception:
        return ""


def load_guidelines(file) -> str:
    if not file:
        return ""
    try:
        content = file.read()
        try:
            return content.decode("utf-8").strip()
        except Exception:
            return content.decode("latin-1", errors="ignore").strip()
    except Exception:
        return ""


def _prefill_last_year_from_prompts(excel_file, context: str) -> Optional[str]:
    """
    Uses ask_variable_list(context) to produce rows and writes them into an excel workbook
    (creates workbook if necessary). Adds a 'change_type' column to help classify rows.
    Returns path to the prefilled workbook or None.
    """
    try:
        try:
            wb = openpyxl.load_workbook(excel_file)
            ws = wb.active
        except Exception:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "variables"

        required_cols = ["variable_name", "old_value", "prompt", "new_value", "change_type"]

        # Ensure header includes change_type
        current_header = [ws.cell(row=1, column=i).value for i in range(1, len(required_cols)+1)]
        if current_header != required_cols:
            for i, col in enumerate(required_cols, start=1):
                ws.cell(row=1, column=i, value=col)

        existing_old_to_row = {}
        for r in range(2, ws.max_row + 1):
            try:
                ov = ws.cell(row=r, column=2).value
                if ov is not None:
                    existing_old_to_row[str(ov).strip()] = r
            except Exception:
                continue

        # Get rows from LLM
        rows = ask_variable_list(context, 4)
        if rows is None:
            return None

        # Insert or update rows from LLM
        for r in rows:
            variable_name = getattr(r, "variable_name", "") or ""
            old_value = (getattr(r, "old_value", "") or "").strip()
            prompt = getattr(r, "prompt", "") or ""
            new_value = getattr(r, "new_value", "") or ""

            if not old_value and not variable_name:
                continue

            if old_value in existing_old_to_row:
                row_idx = existing_old_to_row[old_value]
                ws.cell(row=row_idx, column=1, value=variable_name)
                ws.cell(row=row_idx, column=2, value=old_value)
                if prompt:
                    ws.cell(row=row_idx, column=3, value=prompt)
                if new_value:
                    ws.cell(row=row_idx, column=4, value=new_value)
            else:
                ws.append([variable_name, old_value, prompt, new_value, ""])
                existing_old_to_row[old_value] = ws.max_row

        # Augment with regex-based standard/context detections (do not duplicate existing old_values)
        detected = _detect_standard_and_context_spans(context)
        for old_val, change_type, variable_name in detected:
            if old_val and old_val not in existing_old_to_row:
                ws.append([variable_name or "", old_val, "", "", change_type])
                existing_old_to_row[old_val] = ws.max_row

        tmpdir = tempfile.mkdtemp()
        out_path = os.path.join(tmpdir, "variables_prefilled.xlsx")
        wb.save(out_path)
        return out_path

    except Exception as e:
        print(f"Error in _prefill_last_year_from_prompts: {e}")
        return None


def load_and_annotate_replacements(excel_file, context: str) -> Tuple[Dict[str, str], Optional[str]]:
    """
    Loads/creates workbook; for each row with a prompt calls LLM to fill new_value.
    Returns (replacements_dict, filled_excel_path).
    """
    try:
        wb = openpyxl.load_workbook(excel_file)
        ws = wb.active
    except Exception:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "variables"
        ws.append(["variable_name", "old_value", "prompt", "new_value", "change_type"])

    try:
        headers = [str(c.value).strip().lower() if c.value is not None else "" for c in ws[1]]
    except Exception:
        headers = []

    def _find_col(name: str, default_idx: int) -> int:
        try:
            idx = headers.index(name)
            return idx + 1
        except ValueError:
            return default_idx

    col_placeholder = _find_col("variable_name", 1)
    col_old        = _find_col("old_value", 2)
    col_prompt     = _find_col("prompt", 3)
    col_new        = _find_col("new_value", 4)
    col_type       = _find_col("change_type", 5)

    replacements: Dict[str, str] = {}

    for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
        r = row[0].row

        def _get(col_idx: int) -> str:
            try:
                val = ws.cell(row=r, column=col_idx).value
                return str(val).strip() if val is not None else ""
            except Exception:
                return ""

        variable_name = _get(col_placeholder)
        old_value        = _get(col_old)
        prompt_text      = _get(col_prompt)
        new_value_curr   = _get(col_new)
        change_type_curr = _get(col_type)

        if prompt_text:
            try:
                llm_out = fill_excel_prompts(prompt_text, context, old_value, variable_name)
                llm_out = (llm_out or "").strip()
                if llm_out:
                    ws.cell(row=r, column=col_new, value=llm_out)
                    new_value_curr = llm_out
            except Exception:
                pass

        # If change_type blank, auto-classify based on patterns
        if not change_type_curr:
            detected = _classify_change_type(old_value)
            if detected:
                try:
                    ws.cell(row=r, column=col_type, value=detected)
                    change_type_curr = detected
                except Exception:
                    pass

        if old_value and new_value_curr and old_value != new_value_curr:
            if old_value not in replacements:
                replacements[old_value] = new_value_curr

    # Save annotated copy
    filled_excel_path = None
    try:
        tmpdir = tempfile.mkdtemp()
        filled_excel_path = os.path.join(tmpdir, "variables_filled.xlsx")
        wb.save(filled_excel_path)
    except Exception:
        filled_excel_path = None

    return replacements, filled_excel_path


# =========================
# DOCX helpers
# =========================

def replace_in_paragraph(p, replacements: dict):
    if not replacements:
        return

    repl = {str(k): ("" if v is None else str(v)) for k, v in replacements.items()}

    p_elm = p._p
    ns = p_elm.nsmap or {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "xml": "http://www.w3.org/XML/1998/namespace",
    }

    t_nodes = p_elm.findall(".//w:t", namespaces=ns)
    if not t_nodes:
        return

    originals = [(t, t.text or "") for t in t_nodes]
    full = "".join(txt for _, txt in originals)
    if not full:
        return

    keys = [k for k in repl.keys() if k and (k in full)]
    if not keys:
        return

    keys.sort(key=len, reverse=True)
    unique_keys = [k for k in keys if full.count(k) == 1]
    if not unique_keys:
        return

    pattern = re.compile("|".join(re.escape(k) for k in unique_keys))
    new_full, nsubs = pattern.subn(lambda m: repl[m.group(0)], full)

    if nsubs == 0 or new_full == full:
        return

    lengths = [len(txt) for _, txt in originals]
    pos = 0
    n = len(originals)
    for i in range(n):
        t, _oldtxt = originals[i]
        take = lengths[i] if i < n - 1 else max(0, len(new_full) - pos)
        segment = new_full[pos:pos + take] if take >= 0 else ""
        t.text = segment
        pos += lengths[i] if i < n - 1 else len(new_full) - pos

        if segment and (segment[0].isspace() or segment[-1].isspace()):
            t.set(qn("xml:space"), "preserve")


def replace_placeholders_docx(doc: Document, replacements: dict):
    from docx.oxml.ns import qn
    seen = False
    br_tag = qn('w:br')
    for p in doc.paragraphs:
        if not seen:
            for r in p.runs:
                for br in r._element.findall(br_tag):
                    if br.get(qn('w:type')) == 'page':
                        seen = True
                        break
                if seen:
                    break
            if not seen:
                continue
        replace_in_paragraph(p, replacements)

    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, replacements)

    for sec in doc.sections:
        if sec.header:
            for p in sec.header.paragraphs:
                replace_in_paragraph(p, replacements)
        if sec.footer:
            for p in sec.footer.paragraphs:
                replace_in_paragraph(p, replacements)


def replace_first_page_placeholders_docx(doc: Document, replacements: dict):
    from docx.oxml.ns import qn
    seen = False
    br_tag = qn("w:br")
    typ = qn("w:type")
    for p in doc.paragraphs:
        replace_in_paragraph(p, replacements)
        for r in p.runs:
            for child in r._element:
                if child.tag == br_tag and child.get(typ) == "page":
                    seen = True
                    break
            if seen:
                break
        if seen:
            break


def _build_fallback_docx(replacements: dict, context: str) -> str:
    doc = Document()
    doc.add_heading("Transfer Pricing Output (Fallback)", level=1)

    if replacements:
        doc.add_heading("Resolved Placeholders", level=2)
        tbl = doc.add_table(rows=1, cols=2)
        hdr = tbl.rows[0].cells
        hdr[0].text = "Placeholder"
        hdr[1].text = "Value"
        for k, v in replacements.items():
            row = tbl.add_row().cells
            row[0].text = str(k)
            row[1].text = str(v)
    else:
        doc.add_paragraph("No replacements were generated (missing or empty Excel input).")

    if context:
        doc.add_heading("Context (truncated)", level=2)
        doc.add_paragraph(context[:4000])

    out = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(out.name)
    return out.name


def load_template(template_file) -> str:
    if not template_file:
        return ""
    try:
        doc = Document(template_file)
        chunks = []
        seen = set()

        def add(text: str):
            t = (text or "").strip()
            if t and t not in seen:
                chunks.append(t)
                seen.add(t)

        def add_paragraphs(paragraphs):
            for p in paragraphs:
                add(p.text)

        def add_tables(tables):
            for tbl in tables:
                for row in tbl.rows:
                    cells_txt = []
                    for cell in row.cells:
                        ct = (cell.text or "").strip()
                        if ct:
                            cells_txt.append(ct)
                    if cells_txt:
                        add(" | ".join(cells_txt))

        for sec in doc.sections:
            hdr = sec.header
            if hdr:
                add_paragraphs(hdr.paragraphs)
                add_tables(hdr.tables)

        add_paragraphs(doc.paragraphs)
        add_tables(doc.tables)

        for sec in doc.sections:
            ftr = sec.footer
            if ftr:
                add_paragraphs(ftr.paragraphs)
                add_tables(ftr.tables)

        return "\n".join(chunks)
    except Exception as e:
        print(f"Error loading template: {e}")
        return ""


# -------------------------
# New detection utilities
# -------------------------

# Patterns for standardized date/year tokens

# ------------ Patterns ------------

RE_FY          = re.compile(r'\bFY[E]?\s?\d{2,4}(?:/\d{2})?\b', flags=re.IGNORECASE)
RE_YEAR        = re.compile(r'\b(19|20)\d{2}\b')
RE_YR_RANGE    = re.compile(r'\b(19|20)\d{2}\s?[-–/]\s?(19|20)\d{2}\b')
RE_DATE_DMY    = re.compile(r'\b\d{1,2}\s+(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}\b', re.I)
RE_DATE_MDY    = re.compile(r'\b(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2}(st|nd|rd|th)?[,]?\s+\d{4}\b', re.I)
RE_FINYEAR_PH  = re.compile(r'\b(?:Financial|Fiscal)\s+Year(?:\s+Ended)?[^\n]{0,80}\b(19|20)\d{2}\b', re.I)
RE_FOR_YE_PH   = re.compile(r'\b(?:For|Year)\s+the\s+year\s+ended[^\n]{0,50}\b(19|20)\d{2}\b', re.I)

RE_PERCENT     = re.compile(r'[-(]?\s*\d{1,3}(?:[.,]\d+)?\s?%\s*[)]?')
RE_CURRENCY    = re.compile(r'(?:€|EUR|\\$)\\s*[-(]?[0-9]{1,3}(?:[.,]\\d{3})*(?:[.,]\\d+)?[)]?')
RE_NUMBER      = re.compile(r'\\b-?\\d{1,3}(?:[.,]\\d{3})*(?:[.,]\\d+)?\\b')

RE_SECTIONNUM  = re.compile(r'^\\d+(?:\\.\\d+)+$')  # e.g., 6.3, 2.10.1

LABELS = [
    'net turnover','gross profit','cost of raw material','sales to customers','transaction amount',
    'percentage of the sales','percentage','berry ratio','median','upper quartile','lower quartile',
    'minimum','maximum','number of observation','observations','fte','employees','revenue','expenses',
    'amortisation','depreciation','operating result','weighted average','canada','europe','related parties'
]
LABELS_RE = re.compile(r'(' + r'|'.join(re.escape(x) for x in LABELS) + r')', re.I)

def canonicalize_num(txt: str) -> str:
    t = txt.strip().replace('\u00A0', ' ')
    t = re.sub(r'\bEUR\s*', '€ ', t, flags=re.I)  # normalize EUR
    if t.startswith('(') and t.endswith(')'):
        t = '-' + t[1:-1]
    # collapse floating garbage: 316.08800000000002 -> 316.088
    t = re.sub(r'(\d+\.\d{2})\d+', r'\1', t)
    # remove extra spaces before %
    t = re.sub(r'\s+%', '%', t)
    return t

@dataclass
class Hit:
    source: str
    location: str
    label: str
    value: str
    change_type: str
    confidence: float
    rationale: str

def iter_blocks(doc: Document):
    def _items(parent):
        parent_elm = parent.element if hasattr(parent, 'element') else parent
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield ('p', Paragraph(child, parent))
            elif isinstance(child, CT_Tbl):
                yield ('t', Table(child, parent))
    for block in _items(doc):
        yield block

def sentence_splits(text: str) -> List[str]:
    return re.split(r'(?<=[.!?])\s+', text)

def detect_in_paragraph(text: str) -> List[Hit]:
    hits = []
    # Standardized tokens
    for rx in (RE_FY, RE_DATE_DMY, RE_DATE_MDY, RE_FINYEAR_PH, RE_FOR_YE_PH, RE_YR_RANGE):
        for m in rx.finditer(text):
            hits.append(Hit('DOCX', 'paragraph', 'financial_year', canonicalize_num(m.group(0)),
                           'Standardized change', 0.95, f'match:{rx.pattern[:20]}...'))
    # Contextual numbers
    for sent in sentence_splits(text):
        for m in LABELS_RE.finditer(sent):
            label = m.group(0)
            # Look for % then currency then number in same sentence
            for rx in (RE_PERCENT, RE_CURRENCY, RE_NUMBER):
                nm = rx.search(sent)
                if not nm:
                    continue
                val = canonicalize_num(nm.group(0))
                if RE_SECTIONNUM.match(val):
                    continue
                ctype = 'Standardized change' if RE_YEAR.fullmatch(val) else 'Contextual change'
                conf = 0.80 if rx is RE_CURRENCY else 0.70 if rx is RE_PERCENT else 0.60
                hits.append(Hit('DOCX', 'paragraph', label, val, ctype, conf,
                                'same-sentence label+number'))
                break
    return hits
def detect_in_table(tbl: Table) -> List[Hit]:
    hits = []
    rows = tbl.rows
    for r in rows:
        cells = [c.text.strip() for c in r.cells]
        if not any(cells):
            continue
        # row-wise: if any cell is a label, scan the other cells for numbers
        labels = [c for c in cells if LABELS_RE.search(c)]
        if not labels:
            continue
        numbers = []
        for c in cells:
            for rx in (RE_PERCENT, RE_CURRENCY, RE_NUMBER):
                for m in rx.finditer(c):
                    val = canonicalize_num(m.group(0))
                    if RE_SECTIONNUM.match(val):
                        continue
                    numbers.append(val)
        for lab in labels:
            for val in numbers:
                ctype = 'Standardized change' if RE_YEAR.fullmatch(val) else 'Contextual change'
                conf = 0.92 if RE_CURRENCY.search(val) else 0.82 if RE_PERCENT.search(val) else 0.70
                hits.append(Hit('DOCX', 'table-row', lab, val, ctype, conf,
                                'table row label->value'))
    return hits
def scan_docx(docx_path: str) -> List[Hit]:
    doc = Document(docx_path)
    out = []
    for kind, block in iter_blocks(doc):
        if kind == 'p':
            if block.text.strip():
                out.extend(detect_in_paragraph(block.text))
        else:
            out.extend(detect_in_table(block))
    # dedupe
    seen = set()
    dedup = []
    for h in out:
        key = (h.location, h.label.lower(), h.value, h.change_type)
        if key not in seen:
            seen.add(key)
            dedup.append(h)
    return dedup

def scan_variables_xlsx(xlsx_path: str) -> pd.DataFrame:
    df = pd.read_excel(xlsx_path, engine='openpyxl').fillna('')
    cols = [c.strip().lower() for c in df.columns]
    df.columns = cols
    # normalize names
    col_map = {
        'variable': next((c for c in cols if 'variable' in c), None),
        'old value': next((c for c in cols if 'old' in c and 'value' in c), None),
        'new value': next((c for c in cols if 'new' in c and 'value' in c), None),
        'type':      next((c for c in cols if c=='type' or 'type' in c), None),
        'source':    next((c for c in cols if 'source' in c), None),
    }
    for want, got in col_map.items():
        if got is None:
            raise ValueError(f'Missing column for {want}')
    return df.rename(columns={v:k for k,v in col_map.items() if v})

def build_change_log(doc_hits: List[Hit], df_vars: pd.DataFrame) -> pd.DataFrame:
    # match by value and/or by fuzzy label containment
    rows = []
    for h in doc_hits:
        candidate_type = 'Standardized change' if h.change_type == 'Standardized change' else 'Contextual change'
        rows.append({
            'source': h.source,
            'location': h.location,
            'label_hint': h.label,
            'old_value': h.value,
            'change_type': candidate_type,
            'confidence': h.confidence,
            'rationale': h.rationale,
        })
    out = pd.DataFrame(rows)
    # optional: link to variables sheet by label or value equality
    if not df_vars.empty:
        out['linked_variable'] = ''
        out['proposed_new_value'] = ''
        for i, r in out.iterrows():
            # try exact old value match first
            m = df_vars[df_vars['old value'].astype(str).str.strip() == r['old_value']]
            if m.empty:
                # try label keyword match
                m = df_vars[df_vars['variable'].str.lower().str.contains(re.escape(r['label_hint'].lower()), na=False)]
            if not m.empty:
                out.at[i, 'linked_variable']   = m.iloc[0]['variable']
                out.at[i, 'proposed_new_value'] = m.iloc[0]['new value']
    return out.sort_values(['change_type','confidence'], ascending=[True, False])

def rollforward_year_tokens(text: str, from_year: int, to_year: int) -> str:
    # FY tokens
    text = re.sub(fr'\\bFYE?\\s?{from_year}\\b', f'FY{to_year}', text, flags=re.I)
    text = re.sub(fr'\\bFY\\s?{from_year}\\b',   f'FY{to_year}', text, flags=re.I)
    # full dates
    text = re.sub(fr'(\\b\\d{{1,2}}\\s+(January|February|March|April|May|June|July|August|September|October|November|December)\\s+){from_year}\\b',
                  rf'\\g<1>{to_year}', text, flags=re.I)
    text = re.sub(fr'((January|February|March|April|May|June|July|August|September|October|November|December)\\s+\\d{{1,2}}(?:st|nd|rd|th)?(?:,)?\\s+){from_year}\\b',
                  rf'\\g<1>{to_year}', text, flags=re.I)
    # phrases
    text = re.sub(fr'\\b(Financial|Fiscal)\\s+Year(?:\\s+Ended)?([^\\n]{{0,80}}){from_year}\\b',
                  rf'\\1 Year\\2{to_year}', text, flags=re.I)
    # ranges 2020–2022 -> 2021–2023
    def _bump_range(m):
        a, b = int(m.group(1)), int(m.group(2))
        shift = to_year - from_year
        return f'{a+shift}–{b+shift}'
    text = re.sub(fr'\\b((19|20)\\d{{2})\\s*[–-]\\s*((19|20)\\d{{2})\\b)', 
                  lambda m: f'{int(m.group(2))+to_year-from_year}–{int(m.group(4))+to_year-from_year}', text)
    return text

def _augment_with_regex_detections(path_or_file, context: str) -> Optional[str]:
    """
    Given an excel path or file-like, load workbook, ensure headers include change_type,
    append detected rows from _detect_standard_and_context_spans where the old_value is not already present,
    save to a new temp file and return its path.
    """
    try:
        try:
            wb = openpyxl.load_workbook(path_or_file)
            ws = wb.active
        except Exception:
            # if path_or_file is None or invalid, create blank workbook with headers
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "variables"

        # Ensure header contains change_type column as 5th column
        headers = [ws.cell(row=1, column=i).value for i in range(1, 6)]
        required = ["variable_name", "old_value", "prompt", "new_value", "change_type"]
        for i, name in enumerate(required, start=1):
            if (i > len(headers)) or (headers[i-1] != name):
                ws.cell(row=1, column=i, value=name)

        # Collect existing old_values
        existing = set()
        for r in range(2, ws.max_row + 1):
            try:
                ov = ws.cell(row=r, column=2).value
                if ov is not None:
                    existing.add(str(ov).strip())
            except Exception:
                continue

        detected = _detect_standard_and_context_spans(context)
        for old_val, change_type, hint in detected:
            if old_val and old_val not in existing:
                # Append with placeholder hint (lowercase, underscore)
                placeholder_hint = hint.lower().replace(" ", "_")[:64] if hint else ""
                ws.append([placeholder_hint, old_val, "", "", change_type])
                existing.add(old_val)

        tmpdir = tempfile.mkdtemp()
        out_path = os.path.join(tmpdir, "variables_prefilled_regex.xlsx")
        wb.save(out_path)
        return out_path
    except Exception as e:
        print(f"_augment_with_regex_detections error: {e}")
        return None
