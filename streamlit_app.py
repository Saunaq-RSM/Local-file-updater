import io
import os
import tempfile
import traceback
import pandas as pd
import streamlit as st

# Try to import the real processor. If that fails, capture the traceback and provide
# lightweight fallback implementations so the Streamlit UI still loads and the user
# can see the import error and continue with a minimal flow.
PROCESSOR_OK = False
_processor_import_error = None

try:
    from processor import configure, find_relevant_variables, fill_section_values, generate_doc_from_excel_map
    PROCESSOR_OK = True
except Exception:
    _processor_import_error = traceback.format_exc()
    PROCESSOR_OK = False

    # Minimal fallback "processor" functions so the app can start even if processor.py fails to import.
    # These are intentionally simple and do not depend on heavy third-party packages.
    _CONFIG = {"api_key": None, "api_endpoint": None}

    def configure(api_key: str, api_endpoint: str):
        _CONFIG["api_key"] = api_key
        _CONFIG["api_endpoint"] = api_endpoint

    def _create_blank_variables_xlsx(path: str):
        # Create a minimal variables workbook using pandas (DataFrame -> xlsx).
        df = pd.DataFrame(columns=["variable_name", "old_value", "prompt", "new_value", "change_type"])
        try:
            df.to_excel(path, index=False)
            return True
        except Exception:
            # If writing xlsx fails, fallback to csv (still gives user a file to edit)
            try:
                path_csv = path.rsplit(".", 1)[0] + ".csv"
                df.to_csv(path_csv, index=False)
                return path_csv
            except Exception:
                return False

    def find_relevant_variables(files: dict):
        """
        Fallback implementation:
         - If user uploaded an excel, return it (write to disk if it's an UploadedFile).
         - Otherwise create a blank workbook and a fallback docx that contains the import error.
        Returns (doc_path, filled_excel_path)
        """
        tmpdir = tempfile.mkdtemp()
        out_xlsx = os.path.join(tmpdir, "variables_prefilled_fallback.xlsx")

        # If an excel file was provided as a file-like, persist it and return it
        excel_file = files.get("excel") if files else None
        if excel_file:
            try:
                # streamlit UploadedFile supports read()
                content = excel_file.read()
                with open(out_xlsx, "wb") as f:
                    f.write(content)
                # also produce a simple docx that notes we used the uploaded excel
                doc_path = os.path.join(tmpdir, "fallback_preview.docx")
                try:
                    from docx import Document
                    doc = Document()
                    doc.add_heading("Fallback preview", level=1)
                    doc.add_paragraph("Processor import failed. Using uploaded variables.xlsx as-is.")
                    if _processor_import_error:
                        doc.add_heading("Processor import error (truncated)", level=2)
                        doc.add_paragraph(_processor_import_error[:3000])
                    doc.save(doc_path)
                except Exception:
                    # If python-docx not available, write a plain text fallback
                    doc_path = os.path.join(tmpdir, "fallback_preview.txt")
                    with open(doc_path, "w") as f:
                        f.write("Processor import failed. Using uploaded variables.xlsx as-is.\n\n")
                        if _processor_import_error:
                            f.write("Processor import error:\n")
                            f.write(_processor_import_error)
                return doc_path, out_xlsx
            except Exception:
                pass

        # No excel uploaded — create blank workbook
        created = _create_blank_variables_xlsx(out_xlsx)
        if not created:
            # cannot create xlsx/csv; return None results so UI can show an error
            return None, None

        tmpdir = tempfile.mkdtemp()
        doc_path = os.path.join(tmpdir, "fallback_preview.docx")
        try:
            from docx import Document
            doc = Document()
            doc.add_heading("Fallback preview", level=1)
            doc.add_paragraph("Processor import failed. A blank variables workbook has been created for you to edit.")
            if _processor_import_error:
                doc.add_heading("Processor import error (truncated)", level=2)
                doc.add_paragraph(_processor_import_error[:3000])
            doc.save(doc_path)
        except Exception:
            doc_path = os.path.join(tmpdir, "fallback_preview.txt")
            with open(doc_path, "w") as f:
                f.write("Processor import failed. A blank variables workbook has been created for you to edit.\n\n")
                if _processor_import_error:
                    f.write("Processor import error:\n")
                    f.write(_processor_import_error)

        return doc_path, out_xlsx

    def fill_section_values(files):
        """
        Fallback: simply returns the excel passed in (persist it if it's file-like).
        """
        excel = files.get("excel") if files else None
        if excel and hasattr(excel, "read"):
            tmpdir = tempfile.mkdtemp()
            out_xlsx = os.path.join(tmpdir, "variables_section_filled_fallback.xlsx")
            try:
                with open(out_xlsx, "wb") as f:
                    f.write(excel.read())
                return None, out_xlsx
            except Exception:
                return None
        # If already a path string, return it
        if isinstance(excel, str) and os.path.exists(excel):
            return None, excel
        # otherwise return None so UI can handle it
        return None

    def generate_doc_from_excel_map(file_map, context: str = ""):
        """
        Fallback: read the excel map if present, produce a simple docx listing placeholders/replacements.
        """
        excel = file_map.get("excel")
        tmpdir = tempfile.mkdtemp()
        doc_path = os.path.join(tmpdir, "generated_fallback.docx")

        replacements = {}
        try:
            if excel and isinstance(excel, str) and os.path.exists(excel):
                # try to read using pandas (will work if openpyxl is installed)
                try:
                    df = pd.read_excel(excel, engine="openpyxl")
                    # normalize common names if present
                    cols = [c.lower().strip() for c in df.columns]
                    # Try to identify columns
                    def find(col_names):
                        for c in col_names:
                            if c in cols:
                                return df.columns[cols.index(c)]
                        return None
                    old_col = find(["old_value", "old value", "old"])
                    new_col = find(["new_value", "new value", "new"])
                    if old_col is not None and new_col is not None:
                        for _, row in df.iterrows():
                            o = str(row[old_col]) if not pd.isna(row[old_col]) else ""
                            n = str(row[new_col]) if not pd.isna(row[new_col]) else ""
                            if o and n and o != n:
                                replacements[o] = n
                except Exception:
                    pass
        except Exception:
            pass

        try:
            from docx import Document
            doc = Document()
            doc.add_heading("Generated Document (Fallback)", level=1)
            if replacements:
                doc.add_heading("Applied replacements", level=2)
                tbl = doc.add_table(rows=1, cols=2)
                hdr = tbl.rows[0].cells
                hdr[0].text = "Placeholder"
                hdr[1].text = "Replacement"
                for k, v in replacements.items():
                    row_cells = tbl.add_row().cells
                    row_cells[0].text = str(k)
                    row_cells[1].text = str(v)
            else:
                doc.add_paragraph("No replacements found or processor unavailable.")
            if _processor_import_error:
                doc.add_heading("Processor import error (for debugging)", level=2)
                doc.add_paragraph(_processor_import_error[:4000])
            doc.save(doc_path)
            return doc_path, (excel if isinstance(excel, str) else None)
        except Exception:
            # fallback to plain text
            txt_path = os.path.join(tmpdir, "generated_fallback.txt")
            with open(txt_path, "w") as f:
                if replacements:
                    f.write("Applied replacements:\n")
                    for k, v in replacements.items():
                        f.write(f"{k} -> {v}\n")
                else:
                    f.write("No replacements found or processor unavailable.\n\n")
                if _processor_import_error:
                    f.write("\nProcessor import error:\n")
                    f.write(_processor_import_error)
            return txt_path, (excel if isinstance(excel, str) else None)


# ---- End of dynamic import / fallback setup ----

# Configure backend with secrets (wrapped to avoid KeyError when secrets missing)
try:
    azure_key = st.secrets.get("AZURE_API_KEY")
    azure_endpoint = st.secrets.get("AZURE_API_ENDPOINT")
    if azure_key and azure_endpoint:
        configure(azure_key, azure_endpoint)
except Exception:
    # If configure is not available for some reason, ignore; the fallback configure above will capture calls.
    pass

st.set_page_config(page_title="TP Template Updater", layout="wide")
st.title("TP Agent 2 Yearly Update")

# If processor import failed, show a prominent message with the traceback and suggested fixes
if not PROCESSOR_OK:
    st.error("processor.py failed to import. The app is running in fallback mode.")
    with st.expander("Show processor import traceback (click to expand)"):
        st.code(_processor_import_error or "No traceback captured")
    st.info(
        "Suggested fixes:\n"
        "- Make sure you're running Python 3.10+ (the processor uses modern type syntax).\n"
        "- Ensure required packages are installed: pydantic (v2), python-docx, pdfplumber, openpyxl, requests, pandas.\n"
        "- Run `python -m py_compile processor.py` locally to spot syntax errors (or paste the traceback here for help).\n"
        "While you fix processor.py the app will use a minimal fallback implementation so you can still edit/upload a variables workbook."
    )

st.download_button(
    "Download variables excel",
    data = None if not os.path.exists("documents2/RSM NL - TP Agent 2 Yearly Update Variables - 18.09.2025 V1.xlsx") else open("documents2/RSM NL - TP Agent 2 Yearly Update Variables - 18.09.2025 V1.xlsx", "rb").read(),
    file_name="RSM_NL_Variables_example.xlsx",
    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
)

st.info(
    """
**How this tool works**


1. Upload your files below — Only *Last year local file* is mandatory to upload.

2. Click **Step 1 – Fill & preview variables**. Review/edit the table, Make sure the old values are unique in the text and replacable. Rerunning will rewrite the old table

3. Click **Step 2 – Fill Section Variables**. This will answer any prompts you put in the prompts column of the table from the first step. 

4. Click **Step 3 – Generate final document** to produce the DOCX/PPTX using your edited values.
"""
)

# --- File uploaders (variables + template required) ---
guidelines_file = st.file_uploader("OECD Transfer Pricing Guideline (.txt) — optional", type=["txt"], key="u_guidelines")
transcript_file = st.file_uploader("Client Meeting Transcript (.docx) — optional", type=["docx"], key="u_transcript")
analysis_file   = st.file_uploader("Financial Documents (.pdf) — optional", type=["pdf"], key="u_pdf")
variables_file  = st.file_uploader("Variables (.xlsx) — REQUIRED", type=["xlsx"], key="u_excel")
template_file   = st.file_uploader("Last year local file (.pptx/.docx) — REQUIRED for Step 2", type=["pptx", "docx"], key="u_template")

# Session state: outputs & dataframes
if "generated" not in st.session_state:
    st.session_state.generated = None
if "filled_excel_path" not in st.session_state:
    st.session_state.filled_excel_path = None
if "step2_ready" not in st.session_state:
    st.session_state.step2_ready = False
if "section_excel_path" not in st.session_state:
    st.session_state.section_excel_path = None  # Excel after section-filling step
if "sheet_df" not in st.session_state:
    st.session_state.sheet_df = None
if "sheet_path" not in st.session_state:
    st.session_state.sheet_path = None

# Helper: write DF to a Windows-safe temp file and return the PATH
def _df_to_temp_xlsx_path(df: pd.DataFrame) -> str:
    fd, path = tempfile.mkstemp(suffix=".xlsx")
    os.close(fd)
    df.to_excel(path, index=False)
    return path

def _coerce_df(df: pd.DataFrame) -> pd.DataFrame:
    # Keep it simple and editable in Streamlit
    out = df.copy()
    for col in out.columns:
        out[col] = out[col].fillna("")
    # Make 'prompt' definitely editable
    if "prompt" in out.columns:
        out["prompt"] = out["prompt"].astype("string").fillna("")
    return out

def _set_sheet(df: pd.DataFrame):
    df = _coerce_df(df)
    st.session_state.sheet_df = df
    st.session_state.sheet_path = _df_to_temp_xlsx_path(df)

def _download_current_sheet(label="Download current variables.xlsx", key="dl_current_sheet"):
    if st.session_state.sheet_df is None:
        return
    bio = io.BytesIO()
    st.session_state.sheet_df.to_excel(bio, index=False)
    st.download_button(
        label,
        data=bio.getvalue(),
        file_name="variables_current.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        key=key,
        use_container_width=True,
    )
    
def _persist_editor_from_state():
    if "editor_sheet" in st.session_state:
        _set_sheet(st.session_state.editor_sheet)

# Build the file_map used by the backend; pass None for missing optionals
base_file_map = {
    "guidelines": guidelines_file or None,
    "transcript": transcript_file or None,
    "pdf": analysis_file or None,
    "excel": variables_file or None,
    "template": template_file or None,
}

with st.expander("Selected files"):
    st.write(f"- **guidelines** → {guidelines_file.name if guidelines_file else '(none)'}")
    st.write(f"- **transcript** → {transcript_file.name if transcript_file else '(none)'}")
    st.write(f"- **analysis** → {analysis_file.name if analysis_file else '(none)'}")
    st.write(f"- **variables** → {variables_file.name if variables_file else '(none)'}")
    st.write(f"- **template** → {template_file.name if template_file else '(none)'}")

def run_step1_fill_and_preview():
    with st.spinner("Preparing variables…"):
        try:
            # Clear downstream state so it always reflects the latest Step 1 output
            st.session_state.generated = None
            st.session_state.step2_ready = False
            st.session_state.section_excel_path = None

            file_map_preview = dict(base_file_map)

            result = find_relevant_variables(file_map_preview)
            # find_relevant_variables returns (doc_path, filled_excel_path)
            excel_path = None
            if isinstance(result, tuple):
                excel_path = result[1]
            elif isinstance(result, str):
                excel_path = result

            if not excel_path:
                st.error("No Excel produced by Step 1.")
                return

            # Load the excel and normalize columns to include change_type and consistent headers
            try:
                df = pd.read_excel(excel_path, engine="openpyxl")
            except Exception:
                # If openpyxl missing or reading fails, try pandas default reader (may still fail)
                df = pd.read_excel(excel_path)

            _set_sheet(df)  # <- overwrites session_state.sheet_df and sheet_path
            st.success("Variables filled from the latest run. Edit below.")

        except Exception as e:
            st.error(f"Error in Step 1: {e}")
            # show traceback for debugging
            st.exception(e)

# ---------------- STEP 2 ----------------
def run_step2_fill_sections():
    if st.session_state.sheet_path is None:
        st.error("Please run Step 1 first (or load variables) so we have a current sheet.")
        return

    with st.spinner("Filling section values…"):
        try:
            file_map2 = dict(base_file_map)
            file_map2["excel"] = st.session_state.sheet_path  # pass the current saved path
            result = fill_section_values(file_map2)
            section_excel_path = None
            if isinstance(result, tuple):
                section_excel_path = result[1]
            elif isinstance(result, str):
                section_excel_path = result
            elif result is None:
                section_excel_path = st.session_state.sheet_path

            if not section_excel_path:
                st.error("Section-filling step did not return a valid Excel path.")
                return

            try:
                df2 = pd.read_excel(section_excel_path, engine="openpyxl")
            except Exception:
                df2 = pd.read_excel(section_excel_path)
            _set_sheet(df2)  # update the ONE global sheet
            st.success("Section values filled. You can keep editing below.")

        except Exception as e:
            st.error(f"Error in Step 2: {e}")
            st.exception(e)
# If a sheet exists, we already show the editor above and auto-persist + download button

# ---------------- STEP 3 ----------------
def run_step3_generate():
    if st.session_state.sheet_path is None or st.session_state.sheet_df is None:
        st.error("Please prepare the variables in Step 1 (and optionally Step 2) first.")
        return
    if not template_file:
        st.error("Please upload the Template (.pptx/.docx) for Step 3.")
        return

    with st.spinner("Generating final document…"):
        try:
            file_map3 = dict(base_file_map)
            file_map3["excel"] = st.session_state.sheet_path  # always the latest persisted sheet

            result = generate_doc_from_excel_map(file_map3)
            if isinstance(result, tuple):
                doc_path, excel_path = result
            else:
                doc_path, excel_path = result, None

            # If returned path is not a docx/pptx (text fallback), still show as downloadable plain file
            with open(doc_path, "rb") as f:
                doc_bytes = f.read()

            if doc_path.endswith(".pptx"):
                doc_name = "filled.pptx"
                doc_mime = "application/vnd.openxmlformats-officedocument.presentationml.presentation"
            elif doc_path.endswith(".docx"):
                doc_name = "filled.docx"
                doc_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            else:
                doc_name = os.path.basename(doc_path)
                doc_mime = "text/plain"

            excel_bytes = None
            if excel_path:
                try:
                    with open(excel_path, "rb") as f:
                        excel_bytes = f.read()
                except Exception:
                    excel_bytes = None

            st.session_state.generated = {
                "doc_bytes": doc_bytes,
                "doc_name": doc_name,
                "doc_mime": doc_mime,
                "excel_bytes": excel_bytes,
                "excel_name": "variables_filled.xlsx",
                "excel_mime": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            }
            st.success("Done—download below.")
        except Exception as e:
            st.session_state.generated = None
            st.error(f"Error in Step 3: {e}")
            st.exception(e)

# ---------------- Buttons (inline, no on_click) ----------------
st.divider()
col1, col2, col3 = st.columns(3)
with col1:
    if st.button("Step 1 – Fill & preview variables"):
        run_step1_fill_and_preview()
with col2:
    if st.button("Step 2 – Fill section values"):
        run_step2_fill_sections()
with col3:
    if st.button("Step 3 – Generate final document"):
        run_step3_generate()

if st.session_state.sheet_df is not None:
    st.subheader("Edit variables (live)")

    # Render the editor with the current df
    edited_df = st.data_editor(
        st.session_state.sheet_df,
        key="editor_sheet",          # keep the key if you want, but don't read it from session_state
        use_container_width=True,
        hide_index=True,
        num_rows="dynamic",
    )

    # If the user made a change, persist it and immediately re-run so the editor shows the new base df
    try:
        changed = not st.session_state.sheet_df.equals(edited_df)
    except Exception:
        # equals() can raise if dtypes changed; fall back to a safer check
        changed = True

    if changed:
        _set_sheet(edited_df)   # persists to st.session_state.sheet_df + temp path
        st.rerun()              # refresh now so the UI reflects the new base data this run

# Downloads
gen = st.session_state.generated
if gen:
    st.download_button(
        f"Download {gen['doc_name']}",
        data=gen["doc_bytes"],
        file_name=gen["doc_name"],
        mime=gen["doc_mime"],
        key="dl_doc",
        use_container_width=True,
    )
    if gen["excel_bytes"]:
        st.download_button(
            "Download variables_filled.xlsx",
            data=gen["excel_bytes"],
            file_name=gen["excel_name"],
            mime=gen["excel_mime"],
            key="dl_xlsx",
            use_container_width=True,
        )
